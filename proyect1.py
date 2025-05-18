import streamlit as st
from huggingface_hub import InferenceClient
import requests
from docx import Document
from typing import List
import textwrap

# ---------------- Configuración por defecto ---------------- #

DEFAULT_MODEL_ID = "mistralai/Mistral-7B-Instruct-v0.2"
DEFAULT_HF_TOKEN = st.secrets.get("HF_TOKEN", "")  # Si no se configuró, queda vacío

# ---------------- Cliente Hugging Face ---------------- #

def get_client(token: str | None = None, model_id: str = DEFAULT_MODEL_ID):
    token = token or DEFAULT_HF_TOKEN
    if not token:
        return None  # Modo simple
    return InferenceClient(model=model_id, token=token)

def hf_generate(client, prompt: str, max_tokens: int = 1024, temperature: float = 0.3) -> str:
    if client is None:  # Modo simple
        return simple_llm(prompt)
    return client.text_generation(
        prompt,
        max_new_tokens=max_tokens,
        temperature=temperature,
        top_p=0.9,
        repetition_penalty=1.1,
    )

# ---------------- Modo simple (fallback) ---------------- #

def simple_paragraph(text: str) -> str:
    wrapped = textwrap.fill(text, width=100)
    return wrapped

def simple_llm(prompt: str) -> str:
    placeholders = [
        "Se resaltará la trascendencia del fenómeno planteado y la pertinencia científica de abordarlo.",
        "Se describirá la magnitud del problema a escala global, regional y nacional, destacando sus repercusiones sanitarias y económicas.",
        "Se sintetizarán las principales brechas de conocimiento detectadas en la literatura, subrayando la necesidad de nuevos estudios.",
        "La investigación contribuirá al Objetivo de Desarrollo Sostenible relacionado con la salud y el bienestar.",
        "¿En qué medida el fenómeno descrito se asociará con las variables seleccionadas en el contexto propuesto?",
        "Desde el punto de vista teórico, el estudio profundizará en los modelos conceptuales vigentes y propondrá nuevas perspectivas.",
        "En el plano práctico, los hallazgos orientarán intervenciones basadas en evidencia para mejorar la situación analizada.",
        "Metodológicamente, se empleará un diseño robusto que garantizará la validez y confiabilidad de los resultados obtenidos.",
        "El beneficio social radicará en la generación de conocimiento aplicable que favorecerá la calidad de vida de la población implicada."
    ]
    return "\n\n".join(simple_paragraph(p) for p in placeholders)

# ---------------- Funciones principales ---------------- #

def generate_introduction(client, title: str, objective: str, word_target: int = 4500) -> str:
    prompt = f"""
Redacta una introducción de tesis académica con las siguientes características:
• Extensión mínima: {word_target} palabras (≈ 9 páginas A4).
• Prosa, sin subtítulos, tercera persona, tiempo futuro del modo indicativo.
• Máximo 10 líneas por párrafo.
• Secuencia de párrafos exacta:
  1‑2 p Importancia + plausibilidad del problema.
  1‑2 p Impacto (mundo, Latinoamérica, Perú).
  1‑2 p Vacíos de literatura.
  1 p Contribución a ODS.
  1 p Pregunta de investigación (interrogativa).
  1 p Justificación teórica.
  1 p Justificación práctica.
  1 p Justificación metodológica.
  1 p Justificación social.
Al final escribe la línea EXACTA "===ANTECEDENTES===".
Título: {title}
Objetivo general: {objective}
"""
    return hf_generate(client, prompt, max_tokens=4096)

def paraphrase_abstract(client, abstract: str, word_count: int = 130) -> str:
    if client is None:
        return simple_paragraph("Resumen disponible para consulta; se presentará de forma sintética y libre de plagio en la versión final.")
    prompt = (
        f"Parafrasea en español académico el siguiente resumen en aproximadamente {word_count} palabras, evitando plagio:\n{abstract}"
    )
    return hf_generate(client, prompt, max_tokens=256, temperature=0.4)

def search_pubmed(query: str, n: int = 10) -> List[dict]:
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    ids = requests.get(base + "esearch.fcgi", params={"db": "pubmed", "term": query, "retmax": n, "retmode": "json"}).json()["esearchresult"]["idlist"]
    items = []
    for pmid in ids:
        xml = requests.get(base + "efetch.fcgi", params={"db": "pubmed", "id": pmid, "retmode": "xml"}).text
        import re, html
        abstract = re.search(r"<AbstractText.*?>(.*?)</AbstractText>", xml, re.S)
        if not abstract:
            continue
        authors = re.findall(r"<LastName>(.*?)</LastName>.*?<Initials>(.*?)</Initials>", xml, re.S)
        year = re.search(r"<PubDate>.*?<Year>(\d{4})</Year>", xml, re.S)
        if authors:
            first = f"{authors[0][0]}, {authors[0][1][0]}."
            cite = first + (" et al." if len(authors) > 3 else "")
        else:
            cite = "Autor desconocido"
        items.append({
            "cite": f"{cite} ({year.group(1) if year else 's.f.'})",
            "abstract": html.unescape(abstract.group(1))
        })
    return items[:n]

def build_antecedents(client, pubs: List[dict]) -> str:
    return "\n\n".join(
        f"{p['cite']} {paraphrase_abstract(client, p['abstract'])}" for p in pubs
    ) or "Sin referencias disponibles."

def generate_theoretical_bases(client, title: str, objective: str) -> str:
    if client is None:
        return simple_paragraph("Se desarrollarán los fundamentos conceptuales que sustentan la relación entre las variables propuestas y se explicará el marco teórico que guiará el análisis.")
    prompt = (
        "Redacta las bases teóricas pertinentes a la pregunta de investigación, incluyendo enfoques conceptuales y variables clave. "
        "Prosa académica, tercera persona, futuro indicativo.\n"
        f"Título: {title}\nObjetivo general: {objective}"
    )
    return hf_generate(client, prompt, max_tokens=1024)

def generate_hypotheses(client, objective: str) -> str:
    if client is None:
        return simple_paragraph("Hipótesis de investigación y estadísticas serán formuladas relacionando las variables primarias para demostrar la dirección y fuerza del efecto esperado.")
    prompt = (
        "Formula la hipótesis de investigación y las hipótesis estadísticas (nula y alternativa) basadas en el siguiente objetivo general.\n"
        f"{objective}"
    )
    return hf_generate(client, prompt, max_tokens=256)

def build_docx(intro: str, antecedentes: str, bases: str, hyps: str) -> bytes:
    doc = Document()
    doc.add_heading("Proyecto de Tesis – Secciones Generadas", 1)
    doc.add_heading("Introducción", 2)
    for p in intro.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Antecedentes", 2)
    for p in antecedentes.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Bases Teóricas", 2)
    for p in bases.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Hipótesis", 2)
    for p in hyps.split("\n"):
        doc.add_paragraph(p)

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ---------------- Interfaz Streamlit ---------------- #

st.set_page_config(page_title="Generador de Introducciones de Tesis", layout="wide")
st.title("📝 Generador Automático de Introducciones de Tesis – Modo LLM o Simple")

with st.sidebar:
    st.header("Configuración")
    hf_token = st.text_input("Hugging Face API Token (opcional)", type="password", value="")
    model_id = st.text_input("ID del modelo (HF Hub)", value=DEFAULT_MODEL_ID)
    title_input = st.text_input("Título de la Investigación")
    objective_input = st.text_area("Objetivo General")
    generate_btn = st.button("Generar Introducción")

if generate_btn:
    client = get_client(hf_token, model_id)

    with st.spinner("Generando introducción…"):
        intro = generate_introduction(client, title_input, objective_input)
    st.subheader("Introducción")
    st.markdown(intro)

    with st.spinner("Buscando antecedentes en PubMed…"):
        pubs = search_pubmed(title_input or objective_input) if (title_input or objective_input) else []
        antecedentes = build_antecedents(client, pubs)
    st.subheader("Antecedentes")
    st.markdown(antecedentes)

    with st.spinner("Generando bases teóricas…"):
        bases = generate_theoretical_bases(client, title_input, objective_input)
    st.subheader("Bases Teóricas")
    st.markdown(bases)

    with st.spinner("Generando hipótesis…"):
        hyps = generate_hypotheses(client, objective_input)
    st.subheader("Hipótesis")
    st.markdown(hyps)

    st.success("¡Secciones generadas! Puedes copiar el texto o exportar a Word.")

    docx_bytes = build_docx(intro, antecedentes, bases, hyps)
    st.download_button(
        label="Descargar Word (.docx)",
        data=docx_bytes,
        file_name="proyecto_tesis_generado.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
