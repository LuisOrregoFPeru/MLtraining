import streamlit as st
import textwrap
from docx import Document
from io import BytesIO
from datetime import date

# ---------------------------------------------------------
# Generador de Introducciones de Tesis – Versión Local
# ---------------------------------------------------------
# • No requiere claves API ni acceso a modelos externos.
# • Crea Introducción (≥ 9 págs. A4), Bases Teóricas e Hipótesis.
# • Exporta las secciones a un archivo .docx.
# ---------------------------------------------------------

# ⚠️ set_page_config debe ser la PRIMERA llamada Streamlit
st.set_page_config(page_title="Generador de Introducciones de Tesis", layout="wide")

# --------- Utilidades de formato ---------
LINE_WIDTH = 100  # caracteres ≈ 10‑12 palabras por línea
PARA_WORDS = 90   # palabras objetivo por párrafo

def wrap(text: str) -> str:
    return textwrap.fill(text, width=LINE_WIDTH)

def build_paragraph(core: str) -> str:
    """Garantiza párrafos ≥ PARA_WORDS palabras. Mantiene ≤ 10 líneas aprox."""
    base = core.strip()
    while len(base.split()) < PARA_WORDS:
        base += " " + core.strip()
    return wrap(base)

# --------- Generación de Introducción ---------
def generate_introduction(title: str, objective: str, min_words: int = 4500) -> str:
    title_l = title.lower().strip()
    # Bloques base
    blocks = [
        f"El estudio abordará la problemática relacionada con {title_l}, considerándose un asunto prioritario para la salud pública y la gestión económica.",
        "El interés científico se justificará por la plausibilidad de la pregunta de investigación, la cual abrirá la posibilidad de diseñar intervenciones basadas en evidencia.",
        "Las estadísticas más recientes mostrarán una prevalencia creciente y una morbimortalidad elevada asociada al fenómeno, superando umbrales de alerta en varios continentes.",
        "En América Latina, y particularmente en el Perú, dichos indicadores reflejarán disparidades marcadas entre regiones y grupos socioeconómicos, impactando la sostenibilidad de los sistemas sanitarios y las economías familiares.",
        "La revisión de la literatura revelará una escasez de estudios con diseños robustos y muestras representativas, así como la predominancia de investigaciones en contextos poco comparables.",
        "Persistirán vacíos metodológicos y ausencia de modelos de análisis integrales que contemplen las variables económicas y socioculturales vinculadas al problema.",
        "La investigación contribuirá al Objetivo de Desarrollo Sostenible 3, meta 3.8, orientada a garantizar una vida sana y promover el bienestar para todos en todas las edades.",
        f"¿En qué medida {title_l} se asociará con los resultados planteados en el objetivo general propuesto?",
        "Teóricamente, el trabajo ampliará los marcos de referencia actuales al integrar perspectivas epidemiológicas, económicas y de comportamiento, fortaleciendo la explicación causal del fenómeno.",
        "En el plano práctico, los hallazgos permitirán optimizar estrategias de prevención y asignación de recursos, beneficiando la toma de decisiones de gestores y clínicos.",
        "Metodológicamente, se empleará un diseño de investigación riguroso que garantizará validez interna y externa, con mediciones estandarizadas y análisis multivariados.",
        "Desde la perspectiva social, la generación de conocimiento fomentará la igualdad de oportunidades y reforzará la cohesión comunitaria al reducir las brechas identificadas."
    ]
    paragraphs = [build_paragraph(b) for b in blocks]
    intro = "\n\n".join(paragraphs)
    # Rellenar hasta min_words
    i = 0
    while len(intro.split()) < min_words:
        intro += "\n\n" + paragraphs[i % len(paragraphs)]
        i += 1
    return intro

# --------- Bases Teóricas ---------
def generate_theoretical_bases(objective: str) -> str:
    core = (
        f"El marco teórico contextualizará {objective.lower()}, articulando conceptos derivados del Modelo Socio‑Ecológico, la Teoría del Comportamiento Planificado y la Economía de la Salud. "
        "Se establecerá la variable independiente como determinante principal y la variable dependiente como resultado medible, mientras que factores de confusión actuarán como covariables. "
        "La integración de estos enfoques permitirá explicar las rutas causales, fundamentar la selección de indicadores y definir supuestos para los análisis estadísticos."
    )
    return wrap(core)

# --------- Hipótesis ---------
def generate_hypotheses(title: str) -> str:
    hip_inv = f"Hipótesis de investigación: la presencia de {title.lower()} ejercerá un efecto significativo sobre la variable dependiente planteada en el objetivo general."
    hip_est = (
        "Hipótesis nula (H0): β1 = 0 — no existirá asociación entre la variable independiente y la dependiente.\n"
        "Hipótesis alternativa (H1): β1 ≠ 0 — existirá una asociación estadísticamente significativa entre ambas variables."
    )
    return wrap(hip_inv) + "\n\n" + wrap(hip_est)

# --------- DOCX builder corregido ---------
def build_docx(intro: str, bases: str, hyps: str) -> bytes:
    doc = Document()
    doc.add_heading("Proyecto de Tesis – Secciones Generadas", level=1)

    # Introducción
    doc.add_heading("Introducción", level=2)
    for line in intro.splitlines():
        doc.add_paragraph(line)
    doc.add_page_break()

    # Bases Teóricas
    doc.add_heading("Bases Teóricas", level=2)
    for line in bases.splitlines():
        doc.add_paragraph(line)
    doc.add_page_break()

    # Hipótesis
    doc.add_heading("Hipótesis", level=2)
    for line in hyps.splitlines():
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --------- Interfaz Streamlit ---------
st.title("📝 Generador Automático de Introducciones de Tesis (local)")
with st.sidebar:
    st.header("Entradas")
    title_input = st.text_input("Título de la Investigación")
    objective_input = st.text_area("Objetivo General")

if st.sidebar.button("Generar Secciones"):
    intro_text = generate_introduction(title_input, objective_input)
    bases_text = generate_theoretical_bases(objective_input)
    hyps_text = generate_hypotheses(title_input)

    # Mostrar en pantalla
    st.subheader("Introducción")
    st.markdown(intro_text)

    st.subheader("Bases Teóricas")
    st.markdown(bases_text)

    st.subheader("Hipótesis")
    st.markdown(hyps_text)

    # Botón de descarga
    docx_data = build_docx(intro_text, bases_text, hyps_text)
    st.download_button(
        "Descargar DOCX",
        data=docx_data,
        file_name=f"secciones_tesis_{date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


